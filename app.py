import sys
import os

try:
    __import__('pysqlite3')
    sys.modules['sqlite3'] = sys.modules.pop('pysqlite3')
except ImportError:
    pass

import time
import streamlit as st
import base64
import random
import re
from openai import OpenAI
from dotenv import load_dotenv
from docx import Document
from io import BytesIO
import json

load_dotenv()

try:
    groq_client = OpenAI(
        api_key=os.environ["GROQ_API_KEY"],
        base_url="https://api.groq.com/openai/v1"
    )
except KeyError:
    st.error("GROQ_API_KEY not found. Please set it in your .env file or environment variables.")
    st.stop()

if 'token_usage' not in st.session_state:
    st.session_state.token_usage = {
        'timestamp': time.time(),
        'used': 0
    }

def handle_rate_limit(max_retries=5, initial_delay=20):
    """Decorator to handle rate limit errors with exponential backoff and jitter"""
    def decorator(func):
        def wrapper(*args, **kwargs):
            retries = 0
            delay = initial_delay
            
            current_time = time.time()
            if current_time - st.session_state.token_usage['timestamp'] >= 60:
                st.session_state.token_usage = {
                    'timestamp': current_time,
                    'used': 0
                }
            
            if st.session_state.token_usage['used'] > 5000:
                wait_time = 60 - (current_time - st.session_state.token_usage['timestamp'])
                if wait_time > 0:
                    st.info(f"⏳ Approaching rate limit. Waiting {wait_time:.1f} seconds...")
                    time.sleep(wait_time)
                    st.session_state.token_usage = {
                        'timestamp': time.time(),
                        'used': 0
                    }
            
            while retries < max_retries:
                try:
                    result = func(*args, **kwargs)
                    if hasattr(result, 'usage') and hasattr(result.usage, 'total_tokens'):
                        st.session_state.token_usage['used'] += result.usage.total_tokens
                    return result
                except Exception as e:
                    error_message = str(e)
                    if "rate_limit_exceeded" in error_message or "Rate limit" in error_message or "429" in error_message:
                        retries += 1
                        if retries == max_retries:
                            st.error("Rate limit exceeded after multiple retries. Please try again later.")
                            raise
                        
                        jitter = random.uniform(0.8, 1.2)
                        actual_delay = delay * jitter
                        st.warning(f"⏳ Rate limit reached. Retrying in {actual_delay:.1f} seconds... (Attempt {retries}/{max_retries})")
                        time.sleep(actual_delay)
                        delay *= 2
                        
                        if actual_delay > 10:
                            st.session_state.token_usage = {
                                'timestamp': time.time(),
                                'used': 0
                            }
                    else:
                        raise
            return wrapper
        return wrapper
    return decorator

@handle_rate_limit()
def transcribe_audio(audio_file):
    return groq_client.audio.transcriptions.create(
        model="whisper-large-v3",
        file=audio_file,
        response_format="text"
    )

@handle_rate_limit()
def generate_chat_completion(messages, model="llama-3.1-8b-instant", max_tokens=150):
    return groq_client.chat.completions.create(
        model=model,
        messages=messages,
        temperature=0.3,
        max_tokens=max_tokens
    )

def is_aviation_related(text):
    aviation_keywords = [
        'aircraft', 'airplane', 'helicopter', 'flight', 'pilot', 'ATC', 
        'air traffic', 'runway', 'tower', 'approach', 'landing', 'takeoff',
        'emergency', 'mayday', 'squawk', 'altitude', 'heading', 'vectors',
        'cleared', 'taxi', 'hold short', 'ILS', 'VFR', 'IFR', 'transponder',
        'Cessna', 'Boeing', 'Airbus', 'roger', 'affirm', 'negative', 'departure'
    ]
    text_lower = text.lower()
    return any(keyword in text_lower for keyword in aviation_keywords)

def audio_to_base64(file):
    try:
        with open(file, "rb") as audio_file:
            audio_bytes = audio_file.read()
            base64_audio = base64.b64encode(audio_bytes).decode()
        return base64_audio
    except FileNotFoundError:
        return ""

def extract_section(text, start_marker, end_marker=None):
    if not start_marker:
        return ""
        
    start_marker = start_marker.lower().strip()
    text_lower = text.lower()
    start_idx = text_lower.find(start_marker)
    if start_idx == -1:
        return ""
        
    start_idx_content = text.find(start_marker, start_idx) + len(start_marker)
    if start_idx_content >= len(text):
        return ""

    if end_marker:
        end_marker_lower = end_marker.lower().strip()
        end_idx = text_lower.find(end_marker_lower, start_idx_content)
        return text[start_idx_content:end_idx].strip() if end_idx != -1 else text[start_idx_content:].strip()
    return text[start_idx_content:].strip()

def clean_text(text):
    text = text.replace('**', '').replace('*', '').replace('#', '')
    text = '\n'.join([line.strip() for line in text.split('\n') if line.strip()])
    return text

def extract_report_data(report_text):
    cleaned_report_text = clean_text(report_text)
    
    markers = {
        'date': ('Date:', 'Aircraft:'),
        'aircraft': ('Aircraft:', 'Location:'),
        'location': ('Location:', 'Incident Type:'),
        'incident_type': ('Incident Type:', 'Key Details:'),
        'key_details': ('Key Details:', 'Safety Concerns:'),
        'safety_concerns': ('Safety Concerns:', 'Recommendations:'),
        'recommendations': ('Recommendations:', None)
    }

    data = {
        'date': extract_section(cleaned_report_text, markers['date'][0], markers['date'][1]),
        'aircraft': extract_section(cleaned_report_text, markers['aircraft'][0], markers['aircraft'][1]),
        'location': extract_section(cleaned_report_text, markers['location'][0], markers['location'][1]),
        'incident_type': extract_section(cleaned_report_text, markers['incident_type'][0], markers['incident_type'][1]),
        'key_details': [],
        'safety_concerns': [],
        'recommendations': []
    }
    
    key_details_text = extract_section(cleaned_report_text, markers['key_details'][0], markers['key_details'][1])
    safety_concerns_text = extract_section(cleaned_report_text, markers['safety_concerns'][0], markers['safety_concerns'][1])
    recommendations_text = extract_section(cleaned_report_text, markers['recommendations'][0], markers['recommendations'][1])
    
    for section_text, target in [
        (key_details_text, data['key_details']),
        (safety_concerns_text, data['safety_concerns']),
        (recommendations_text, data['recommendations'])
    ]:
        if not section_text:
            continue
        for line in section_text.split('\n'):
            line = line.strip()
            if line:
                clean_line = re.sub(r'^(\*|-|•)\s*', '', line).strip()
                if clean_line:
                    target.append(clean_line)
    
    data['date'] = data['date'] or 'Not specified'
    data['aircraft'] = data['aircraft'] or 'Not specified'
    data['location'] = data['location'] or 'Not specified'
    data['incident_type'] = data['incident_type'] or 'Not specified'
    
    return data

def generate_markdown_report(report_data):
    markdown = "# 📋 Flight Incident Investigation Report\n\n"
    
    markdown += "## 📊 General Information\n\n"
    markdown += f"**📅 Date:** {report_data.get('date', 'Not specified')}  \n"
    markdown += f"**✈️ Aircraft:** {report_data.get('aircraft', 'Not specified')}  \n"
    markdown += f"**📍 Location:** {report_data.get('location', 'Not specified')}  \n"
    markdown += f"**⚠️ Incident Type:** {report_data.get('incident_type', 'Not specified')}\n\n"
    
    markdown += "---\n\n"
    
    markdown += "## 🔍 Key Details\n\n"
    details = report_data.get('key_details', [])
    if not details:
        markdown += "* No key details were extracted.\n"
    for detail in details:
        markdown += f"* {detail}\n"
    markdown += "\n"
    
    markdown += "## 🚨 Safety Concerns\n\n"
    concerns = report_data.get('safety_concerns', [])
    if not concerns:
        markdown += "* No specific safety concerns were identified.\n"
    for concern in concerns:
        markdown += f"* {concern}\n"
    markdown += "\n"
    
    markdown += "## 💡 Recommendations\n\n"
    recommendations = report_data.get('recommendations', [])
    if not recommendations:
        markdown += "* No specific recommendations were provided.\n"
    for recommendation in recommendations:
        markdown += f"* {recommendation}\n"
        
    return markdown

def generate_docx(report_data):
    doc = Document()
    doc.add_heading('Flight Incident Investigation Report', 0)
    
    doc.add_heading('General Information', level=1)
    
    p = doc.add_paragraph()
    p.add_run('Date: ').bold = True
    p.add_run(report_data.get('date', 'Not specified'))
    
    p = doc.add_paragraph()
    p.add_run('Aircraft: ').bold = True
    p.add_run(report_data.get('aircraft', 'Not specified'))
    
    p = doc.add_paragraph()
    p.add_run('Location: ').bold = True
    p.add_run(report_data.get('location', 'Not specified'))
    
    p = doc.add_paragraph()
    p.add_run('Incident Type: ').bold = True
    p.add_run(report_data.get('incident_type', 'Not specified'))
    
    doc.add_heading('Key Details', level=1)
    for detail in report_data.get('key_details', []):
        doc.add_paragraph(detail, style='List Bullet')
    
    doc.add_heading('Safety Concerns', level=1)
    for concern in report_data.get('safety_concerns', []):
        doc.add_paragraph(concern, style='List Bullet')
    
    doc.add_heading('Recommendations', level=1)
    for recommendation in report_data.get('recommendations', []):
        doc.add_paragraph(recommendation, style='List Bullet')
    
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def generate_incident_report_with_groq(transcript):
    cleaning_prompt = f"""
    Clean and structure this ATC/blackbox transcript. Identify speakers (PILOT, ATC, etc.) and organize the conversation clearly:
    
    {transcript}
    
    Provide a cleaned, well-formatted version with clear speaker labels.
    """
    
    with st.spinner("🔄 Cleaning transcript..."):
        cleaning_response = generate_chat_completion(
            [{"role": "user", "content": cleaning_prompt}],
            model="llama-3.1-8b-instant",
            max_tokens=2000
        )
        cleaned_transcript = cleaning_response.choices[0].message.content
    
    report_prompt = f"""
    Based on this aviation transcript, generate a comprehensive flight incident investigation report.
    
    CLEANED TRANSCRIPT:
    {cleaned_transcript}
    
    Generate a structured report with the following EXACT sections and formatting:
    
    Date: [Extract or infer date from transcript. Use 'Not specified' if unknown]
    Aircraft: [Aircraft type/registration. Use 'Not specified' if unknown]
    Location: [Airport/geographic area. Use 'Not specified' if unknown]
    Incident Type: [Brief description of incident type - e.g., Near-miss, System failure, Communication issue]
    
    Key Details:
    * [First key event or fact]
    * [Second key event or fact] 
    * [Third key event or fact]
    * [Additional key details as needed]
    
    Safety Concerns:
    * [Primary safety concern identified]
    * [Secondary safety concern]
    * [Additional concerns as relevant]
    
    Recommendations:
    * [First actionable recommendation]
    * [Second actionable recommendation]
    * [Third actionable recommendation]
    
    IMPORTANT: 
    - Use asterisks (*) for all bullet points
    - Start directly with "Date:" without any preamble
    - Be factual and concise
    - Focus on aviation safety aspects
    """
    
    with st.spinner("📝 Generating incident report..."):
        report_response = generate_chat_completion(
            [{"role": "user", "content": report_prompt}],
            model="llama-3.1-8b-instant", 
            max_tokens=2000
        )
        report_text = report_response.choices[0].message.content
    
    return report_text

# --- Enhanced Streamlit UI ---

st.set_page_config(
    layout="wide",
    page_title="Black Box Archive Analysis",
    page_icon="✈️",
    initial_sidebar_state="collapsed"
)

# Custom CSS for professional styling
st.markdown("""
<style>
    /* Main Theme Colors */
    :root {
        --primary-color: #1e40af;
        --secondary-color: #3b82f6;
        --accent-color: #60a5fa;
        --danger-color: #dc2626;
        --success-color: #16a34a;
        --warning-color: #f59e0b;
        --background-dark: #0f172a;
        --card-background: #ffffff;
    }
    
    /* Hide Streamlit Branding */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    
    /* Main Container Styling */
    .main {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 0 !important;
    }
    
    .block-container {
        padding: 2rem 3rem !important;
        max-width: 100% !important;
    }
    
    /* Header Styling */
    .main-header {
        background: linear-gradient(135deg, #1e3a8a 0%, #1e40af 100%);
        padding: 2rem 3rem;
        border-radius: 16px;
        margin-bottom: 2rem;
        box-shadow: 0 10px 40px rgba(0,0,0,0.2);
        border: 1px solid rgba(255,255,255,0.1);
    }
    
    .main-header h1 {
        color: white;
        font-size: 2.5rem;
        font-weight: 700;
        margin: 0;
        text-shadow: 2px 2px 4px rgba(0,0,0,0.3);
    }
    
    .main-header p {
        color: rgba(255,255,255,0.9);
        font-size: 1.1rem;
        margin: 0.5rem 0 0 0;
    }
    
    /* Card Styling */
    .custom-card {
        background: white;
        border-radius: 12px;
        padding: 1.5rem;
        box-shadow: 0 4px 6px rgba(0,0,0,0.07);
        border: 1px solid #e5e7eb;
        height: 100%;
        transition: all 0.3s ease;
    }
    
    .custom-card:hover {
        box-shadow: 0 8px 16px rgba(0,0,0,0.12);
        transform: translateY(-2px);
    }
    
    .card-header {
        font-size: 1.3rem;
        font-weight: 600;
        color: #1e40af;
        margin-bottom: 1rem;
        padding-bottom: 0.75rem;
        border-bottom: 2px solid #e5e7eb;
        display: flex;
        align-items: center;
        gap: 0.5rem;
    }
    
    /* Audio Player Styling */
    audio {
        width: 100%;
        border-radius: 8px;
        margin: 1rem 0;
    }
    
    /* Transcript Box */
    .transcript-box {
        background: #f8fafc;
        border: 1px solid #e2e8f0;
        border-radius: 8px;
        padding: 1rem;
        max-height: 200px;
        overflow-y: auto;
        font-family: 'Courier New', monospace;
        font-size: 0.9rem;
        margin: 1rem 0;
    }
    
    /* Button Styling */
    .stButton>button {
        width: 100%;
        background: linear-gradient(135deg, #3b82f6 0%, #1e40af 100%);
        color: white;
        border: none;
        padding: 0.75rem 1.5rem;
        font-size: 1rem;
        font-weight: 600;
        border-radius: 8px;
        transition: all 0.3s ease;
        box-shadow: 0 4px 6px rgba(59, 130, 246, 0.3);
    }
    
    .stButton>button:hover {
        background: linear-gradient(135deg, #2563eb 0%, #1e3a8a 100%);
        box-shadow: 0 6px 12px rgba(59, 130, 246, 0.4);
        transform: translateY(-2px);
    }
    
    /* Chat Interface */
    .chat-container {
        background: white;
        border-radius: 12px;
        overflow: hidden;
        box-shadow: 0 4px 6px rgba(0,0,0,0.07);
        border: 1px solid #e5e7eb;
    }
    
    .chat-messages {
        height: 450px;
        overflow-y: auto;
        padding: 1.5rem;
        background: #f9fafb;
        scrollbar-width: thin;
        scrollbar-color: #cbd5e1 #f1f5f9;
    }
    
    .chat-messages::-webkit-scrollbar {
        width: 8px;
    }
    
    .chat-messages::-webkit-scrollbar-track {
        background: #f1f5f9;
    }
    
    .chat-messages::-webkit-scrollbar-thumb {
        background: #cbd5e1;
        border-radius: 4px;
    }
    
    .user-message {
        background: linear-gradient(135deg, #3b82f6 0%, #2563eb 100%);
        color: white;
        border-radius: 18px 18px 4px 18px;
        padding: 0.75rem 1rem;
        margin: 0.5rem 0 0.5rem auto;
        max-width: 80%;
        width: fit-content;
        box-shadow: 0 2px 4px rgba(59, 130, 246, 0.3);
        animation: slideInRight 0.3s ease;
    }
    
    .assistant-message {
        background: white;
        border: 1px solid #e5e7eb;
        color: #1f2937;
        border-radius: 18px 18px 18px 4px;
        padding: 0.75rem 1rem;
        margin: 0.5rem auto 0.5rem 0;
        max-width: 80%;
        width: fit-content;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        animation: slideInLeft 0.3s ease;
    }
    
    .welcome-message {
        background: linear-gradient(135deg, #f0f9ff 0%, #e0f2fe 100%);
        border: 1px solid #bae6fd;
        border-radius: 12px;
        padding: 1.25rem;
        margin-bottom: 1rem;
        box-shadow: 0 2px 8px rgba(59, 130, 246, 0.1);
    }
    
    .welcome-message b {
        color: #1e40af;
        font-size: 1.1rem;
    }
    
    .welcome-message ul {
        margin: 0.75rem 0 0 0;
        padding-left: 1.5rem;
        color: #475569;
    }
    
    .welcome-message li {
        margin: 0.5rem 0;
    }
    
    /* Chat Input Area */
    .chat-input-container {
        padding: 1rem;
        background: white;
        border-top: 1px solid #e5e7eb;
    }
    
    /* Animations */
    @keyframes slideInRight {
        from {
            opacity: 0;
            transform: translateX(20px);
        }
        to {
            opacity: 1;
            transform: translateX(0);
        }
    }
    
    @keyframes slideInLeft {
        from {
            opacity: 0;
            transform: translateX(-20px);
        }
        to {
            opacity: 1;
            transform: translateX(0);
        }
    }
    
    /* Status Badges */
    .status-badge {
        display: inline-block;
        padding: 0.25rem 0.75rem;
        border-radius: 12px;
        font-size: 0.875rem;
        font-weight: 600;
        margin: 0.25rem;
    }
    
    .badge-success {
        background: #dcfce7;
        color: #166534;
    }
    
    .badge-warning {
        background: #fef3c7;
        color: #92400e;
    }
    
    .badge-info {
        background: #dbeafe;
        color: #1e40af;
    }
    
    /* Report Section Styling */
    .report-section h1 {
        color: #1e40af;
        border-bottom: 3px solid #3b82f6;
        padding-bottom: 0.5rem;
        margin-bottom: 1.5rem;
    }
    
    .report-section h2 {
        color: #1e40af;
        margin-top: 1.5rem;
        font-size: 1.4rem;
    }
    
    .report-section hr {
        border: none;
        border-top: 2px solid #e5e7eb;
        margin: 1.5rem 0;
    }
    
    /* Download Button */
    a[download] {
        display: inline-block;
        background: linear-gradient(135deg, #16a34a 0%, #15803d 100%);
        color: white !important;
        padding: 0.75rem 1.5rem;
        text-decoration: none;
        border-radius: 8px;
        font-weight: 600;
        transition: all 0.3s ease;
        box-shadow: 0 4px 6px rgba(22, 163, 74, 0.3);
        text-align: center;
    }
    
    a[download]:hover {
        background: linear-gradient(135deg, #15803d 0%, #14532d 100%);
        box-shadow: 0 6px 12px rgba(22, 163, 74, 0.4);
        transform: translateY(-2px);
    }
    
    /* File Uploader */
    .stFileUploader {
        border: 2px dashed #cbd5e1;
        border-radius: 8px;
        padding: 1rem;
        background: #f8fafc;
        transition: all 0.3s ease;
    }
    
    .stFileUploader:hover {
        border-color: #3b82f6;
        background: #f0f9ff;
    }
    
    /* Info boxes */
    .stAlert {
        border-radius: 8px;
        border-left: 4px solid;
    }
    
    /* Text Input in Chat */
    .stTextInput>div>div>input {
        border-radius: 8px;
        border: 1px solid #cbd5e1;
        padding: 0.75rem;
    }
    
    .stTextInput>div>div>input:focus {
        border-color: #3b82f6;
        box-shadow: 0 0 0 3px rgba(59, 130, 246, 0.1);
    }
</style>
""", unsafe_allow_html=True)

# Header
st.markdown("""
<div class="main-header">
    <h1>✈️ Black Box Archive Analysis</h1>
    <p>Advanced AI-powered aviation incident analysis using Groq's LLM technology</p>
</div>
""", unsafe_allow_html=True)

# Initialize session state
if 'messages' not in st.session_state:
    st.session_state.messages = []
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []

# Create three columns
left_col, mid_col, right_col = st.columns([1, 2, 1])

# --- Left Column: Upload & Analyze ---
with left_col:
    st.markdown('<div class="custom-card">', unsafe_allow_html=True)
    st.markdown('<div class="card-header">🎙️ Upload & Analyze</div>', unsafe_allow_html=True)
    
    uploaded_file = st.file_uploader(
        "Upload MP3 Audio File", 
        type=["mp3"], 
        key="uploader",
        help="Upload an aviation black box or ATC recording",
        on_change=lambda: [st.session_state.pop(k, None) for k in ['transcript', 'report_text', 'chat_history', 'messages']]
    )

    if uploaded_file is not None:
        temp_file_path = "uploaded_file.mp3"
        with open(temp_file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        st.markdown('<span class="status-badge badge-success">✓ File uploaded successfully</span>', unsafe_allow_html=True)
        
        base64_audio = audio_to_base64(temp_file_path)
        st.markdown("### 🔊 Audio Preview")
        st.markdown(f"""
        <audio controls style="width: 100%; margin: 1rem 0;">
            <source src="data:audio/mp3;base64,{base64_audio}" type="audio/mp3">
            Your browser does not support the audio element.
        </audio>
        """, unsafe_allow_html=True)
        
        if 'transcript' in st.session_state:
            st.markdown("### 📝 Transcript Preview")
            with st.expander("View Transcript", expanded=False):
                st.text_area("", st.session_state.transcript, height=150, label_visibility="collapsed")

        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("🚀 Analyze Black Box Audio", type="primary", key="analyze_btn"):
            try:
                with st.spinner("🎧 Transcribing audio with Groq Whisper..."):
                    with open(temp_file_path, "rb") as audio_file:
                        transcript_result = transcribe_audio(audio_file)
                        transcript = transcript_result.text if hasattr(transcript_result, 'text') else transcript_result

                st.session_state.transcript = transcript
                st.success("✅ Transcription Complete!")
                
                with st.spinner("🔍 Validating content..."):
                    is_aviation = is_aviation_related(transcript)
                
                if not is_aviation:
                    st.error("⚠️ This audio does not appear to be aviation-related.")
                    st.warning("Please upload an audio file with flight or ATC content.")
                    if 'report_text' in st.session_state:
                        del st.session_state.report_text
                    st.stop()
                else:
                    st.success("✅ Aviation content detected!")
                    
                    with st.spinner("📝 Generating Incident Report..."):
                        report_text = generate_incident_report_with_groq(transcript)
                        st.session_state.report_text = report_text
                        st.session_state.messages = [
                            {"role": "system", "content": f"""You are an expert aviation analyst. Answer user questions based ONLY on the details found in this flight incident report: {report_text}"""}
                        ]
                        st.session_state.chat_history = [] 
                        st.success("✅ Report Generation Complete!")
                        # st.balloons()

            except Exception as e:
                st.error(f"❌ An error occurred: {str(e)}")
                if "rate_limit_exceeded" in str(e) or "Rate limit" in str(e):
                    st.info("⏳ Rate limit hit. Please wait a minute and try again.")
    
    st.markdown('</div>', unsafe_allow_html=True)

# --- Middle Column: Generated Report ---
with mid_col:
    st.markdown('<div class="custom-card">', unsafe_allow_html=True)
    st.markdown('<div class="card-header">📋 Investigation Report</div>', unsafe_allow_html=True)
    
    if 'report_text' in st.session_state:
        try:
            report_data = extract_report_data(st.session_state.report_text)
            
            st.markdown('<div class="report-section">', unsafe_allow_html=True)
            st.markdown(generate_markdown_report(report_data))
            st.markdown('</div>', unsafe_allow_html=True)
            
            st.markdown("---")
            st.markdown("### 💾 Export Report")
            doc_bio = generate_docx(report_data)
            
            b64 = base64.b64encode(doc_bio.read()).decode()
            download_link = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="flight_incident_report.docx">📥 Download Report (.docx)</a>'
            st.markdown(download_link, unsafe_allow_html=True)
            
        except Exception as e:
            st.error(f"❌ Error displaying report: {e}")
            st.text_area("Raw Report Output:", st.session_state.report_text, height=400)
    else:
        st.markdown("""
        <div style="text-align: center; padding: 3rem 1rem;">
            <div style="font-size: 4rem; margin-bottom: 1rem;">📄</div>
            <h3 style="color: #64748b; margin-bottom: 0.5rem;">No Report Generated Yet</h3>
            <p style="color: #94a3b8;">Upload an audio file and click 'Analyze Black Box Audio' to generate a comprehensive incident report.</p>
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown('</div>', unsafe_allow_html=True)

# --- Right Column: Report Assistant Chat ---
with right_col:
    st.markdown('<div class="custom-card">', unsafe_allow_html=True)
    st.markdown('<div class="card-header">💬 Report Assistant</div>', unsafe_allow_html=True)
    
    st.markdown('<div class="chat-container">', unsafe_allow_html=True)
    st.markdown('<div class="chat-messages">', unsafe_allow_html=True)

    if 'report_text' in st.session_state and not st.session_state.chat_history:
        st.markdown("""
        <div class="welcome-message">
            <b>👋 Hello! I'm ready to analyze the report.</b><br><br>
            Ask me questions about the incident, such as:
            <ul>
                <li>What was the primary incident type?</li>
                <li>List the safety recommendations.</li>
                <li>What aircraft was involved?</li>
                <li>What were the key safety concerns?</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)
    elif 'report_text' not in st.session_state:
        st.markdown("""
        <div style="text-align: center; padding: 2rem 1rem;">
            <div style="font-size: 3rem; margin-bottom: 1rem;">🤖</div>
            <p style="color: #94a3b8; font-size: 0.95rem;">The Report Assistant will be available after the incident report is generated.</p>
        </div>
        """, unsafe_allow_html=True)

    for msg in st.session_state.chat_history:
        if msg['role'] == 'user':
            st.markdown(f'<div class="user-message">{msg["content"]}</div>', unsafe_allow_html=True)
        else:
            st.markdown(f'<div class="assistant-message">{msg["content"]}</div>', unsafe_allow_html=True)

    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown('<div class="chat-input-container">', unsafe_allow_html=True)
    with st.form(key='chat_form', clear_on_submit=True):
        user_input = st.text_input(
            "Ask about the incident report",
            key="user_query",
            label_visibility="collapsed",
            placeholder="Type your question here..."
        )
        submitted = st.form_submit_button("Send 📤", type="primary", use_container_width=True)
    st.markdown('</div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

    if submitted and user_input.strip():
        if 'report_text' not in st.session_state:
            st.error("⚠️ Please generate a report first.")
        else:
            st.session_state.chat_history.append({"role": "user", "content": user_input})

            try:
                prompt = f"""
                Based ONLY on the following flight incident report:
                ---REPORT START---
                {st.session_state.report_text}
                ---REPORT END---

                Question: {user_input}

                Provide a concise 1-2 sentence answer.
                If the answer cannot be directly inferred from the report, you MUST respond: "This specific detail is not available in the generated report."
                """
                
                with st.spinner("🤔 Assistant thinking..."):
                    response = generate_chat_completion(
                        [{"role": "user", "content": prompt}],
                        max_tokens=100, 
                        model="llama-3.1-8b-instant"
                    )

                answer = response.choices[0].message.content
                st.session_state.chat_history.append({"role": "assistant", "content": answer})
                st.rerun()

            except Exception as e:
                st.error(f"❌ Error generating response: {str(e)}")
                if "rate_limit_exceeded" in str(e) or "Rate limit" in str(e):
                    st.info("⏳ Please wait a moment and try again.")
    
    st.markdown('</div>', unsafe_allow_html=True)

# Footer
st.markdown("<br><br>", unsafe_allow_html=True)
st.markdown("""
<div style="text-align: center; padding: 2rem; color: white; font-size: 0.9rem;">
    <p style="margin: 0;">Powered by <strong>Groq AI</strong> | Whisper-large-v3 & Llama-3.1-8b</p>
    <p style="margin: 0.5rem 0 0 0; opacity: 0.8;">Advanced Aviation Safety Analysis System</p>
</div>
""", unsafe_allow_html=True)
